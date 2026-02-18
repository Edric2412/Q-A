import os
import shutil
import json
import datetime
import re
import logging
import time
import asyncio
import random
from typing import List, Dict, Any, Optional
from pathlib import Path

# --- Third Party Imports ---
import docx
import pdfplumber
import core.database as db_module
import google.generativeai as genai
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer, util
from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
from google.api_core.exceptions import ResourceExhausted, ServiceUnavailable

# --- Setup & Configuration ---
load_dotenv(Path(__file__).parent.parent / ".env")
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define the Sub-App
evaluator_app = FastAPI(title="AI Answer Evaluator", version="3.5")

evaluator_app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Database ---
# PostgreSQL via asyncpg (see database.py)
# The pool is initialized when main.py starts (shared via db_module)

# --- AI Models ---
GEMINI_API_KEY = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")

if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    # Target Gemini 3 Flash Preview as requested
    try:
        gemini_model = genai.GenerativeModel('gemini-3-flash-preview')
        logger.info("SUCCESS: Gemini 3 Flash initialized.")
    except:
        logger.warning("Gemini 3 Flash unavailable. Falling back to 2.5 Flash.")
        gemini_model = genai.GenerativeModel('gemini-2.5-flash')
else:
    logger.error("ERROR: GEMINI_API_KEY not found.")
    gemini_model = None

try:
    minilm_model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
    logger.info("SUCCESS: MiniLM initialized.")
except Exception as e:
    minilm_model = None
    logger.warning(f"MiniLM failed to load: {e}")

# --- STRICT EXAM SCHEMAS ---
# Enforces: CIA (4/10) and Model (5/8)
EXAM_PATTERNS = {
    "CIA": {
        "mcq":   {"count": 10, "start": 1,  "marks": 1.0},
        "short": {"count": 5,  "start": 11, "marks": 4.0},
        "long":  {"count": 2,  "start": 16, "marks": 10.0}
    },
    "Model": {
        "mcq":   {"count": 10, "start": 1,  "marks": 1.0},
        "short": {"count": 5,  "start": 11, "marks": 5.0},
        "long":  {"count": 5,  "start": 16, "marks": 8.0}
    }
}

# --- Pydantic Models ---
class StudentListRequest(BaseModel):
    department: str
    batch: str
    semester: Optional[str] = None
    subject: Optional[str] = None
    exam_type: Optional[str] = None
    exam_id: Optional[str] = None

class MarkUpdateRequest(BaseModel):
    exam_id: str
    roll_no: str
    question_num: str
    new_mark: float

# --- Helper Functions ---


from utils.vision_utils import grade_pdf_with_vision, extract_first_page_text_ocr

def parse_docx_table_data(file_path: str, is_question_paper: bool = False) -> Dict[str, Dict]:
    doc = docx.Document(file_path)
    items = {} 
    valid_q_pattern = re.compile(r'^(\d+)') 

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2: continue
            
            col0_text = cells[0].text.strip()
            if "Q.No" in col0_text or "Answers" in cells[1].text: continue

            q_match = valid_q_pattern.match(col0_text)
            if q_match:
                q_id = q_match.group(1)
                final_rubric_text = ""

                # Answer Key Logic
                if not is_question_paper and len(cells) > 2:
                    ans_cell = cells[1]
                    mark_cell = cells[2]
                    
                    ans_paras = [p.text.strip() for p in ans_cell.paragraphs if p.text.strip()]
                    mark_paras = [p.text.strip() for p in mark_cell.paragraphs if p.text.strip()]
                    
                    # Strategy A: Intelligent Line-by-Line Mapping
                    if len(ans_paras) == len(mark_paras) and len(ans_paras) > 0:
                        mapped_lines = []
                        for txt, mk in zip(ans_paras, mark_paras):
                            mapped_lines.append(f"{txt} [Value: {mk}]")
                        final_rubric_text = "\n".join(mapped_lines)
                    
                    # Strategy B: Fallback (Just dump everything)
                    else:
                        # This ensures we NEVER lose the marks, even if formatting is weird
                        raw_marks = mark_cell.text.strip()
                        final_rubric_text = f"{ans_cell.text.strip()} [Rubric Breakdown: {raw_marks}]"
                
                else:
                    final_rubric_text = cells[1].text.strip()

                if q_id in items:
                    prefix = " OR " if is_question_paper else "\n[OR Rubric]: "
                    items[q_id]['text'] += f"{prefix}{final_rubric_text}"
                else:
                    items[q_id] = {'text': final_rubric_text}

    return items

def extract_text(file_path: str) -> str:
    """
    Extracts text from PDF or DOCX files.
    CRITICAL FIX: Reads text from BOTH paragraphs and tables in DOCX.
    """
    text_content = []
    
    if file_path.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
            
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        
        # 1. Read Standard Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())
                
        # 2. Read Text inside Tables (Fixes empty answers in grid layouts)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    # Extract text from paragraphs within the cell
                    cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    # Join cell content with a pipe | to keep context
                    text_content.append(" | ".join(row_text))
                    
    return "\n".join(text_content)

def extract_student_identity(text: str) -> str:
    match = re.search(r"(?i)Roll[-\s\.]*(?:No\.?|Number|Num)?\s*[:\-\.]*\s*([A-Z0-9]+)", text)
    return match.group(1).upper().strip() if match else None

def parse_student_text(text: str) -> Dict[str, str]:
    """
    Robustly extracts answers using regex finditer to handle
    variable whitespace and formatting (e.g. 11a, 11., Q11).
    """
    # Pattern looks for a line start or whitespace, followed by number, optional letter, and separator
    pattern = re.compile(r'(?:^|\n)\s*(?:Q\.?|Ans\.?|Answer)?\s*(\d+)(?:\s*[a-zA-Z])?\s*[.)\-\:|]')
    
    matches = list(pattern.finditer(text))
    parsed_answers = {}
    
    for i, match in enumerate(matches):
        q_num = match.group(1)
        start_idx = match.end()
        
        # End index is the start of the next match, or end of string
        end_idx = matches[i+1].start() if i + 1 < len(matches) else len(text)
        
        answer_content = text[start_idx:end_idx].strip()
        
        # Append if duplicate (unlikely but safe)
        if q_num in parsed_answers:
            parsed_answers[q_num] += " " + answer_content
        else:
            parsed_answers[q_num] = answer_content
            
    # Fallback for completely unformatted text
    if not parsed_answers:
        paragraphs = [p.strip() for p in text.split('\n') if len(p.strip()) > 5]
        for idx, para in enumerate(paragraphs):
            parsed_answers[str(idx + 1)] = para
            
    return parsed_answers
    
async def call_gemini_api_safe(prompt: str, retries=3):
    """
    Calls Gemini with strict rate limiting and backoff to handle 429 errors.
    """
    base_delay = 10  # Wait 10 seconds between calls (Free tier is ~2-4 RPM safe)
    
    for attempt in range(retries):
        try:
            # 1. Force a wait BEFORE every call to respect Rate Limits
            await asyncio.sleep(base_delay) 
            
            # Async generation
            response = await gemini_model.generate_content_async(prompt)
            return response
            
        except Exception as e:
            error_str = str(e).lower()
            if "429" in error_str or "quota" in error_str:
                wait_time = (attempt + 1) * 20  # Wait 20s, then 40s, then 60s
                logger.warning(f"Quota Hit! Sleeping for {wait_time}s before retry {attempt+1}/{retries}...")
                await asyncio.sleep(wait_time)
            else:
                logger.error(f"Gemini API Error: {e}")
                return None
    return None

def clean_json_string(text: str) -> str:
    """Extracts the first valid JSON object from a string."""
    try:
        start = text.find('{')
        end = text.rfind('}')
        if start != -1 and end != -1:
            return text[start:end+1]
        return text
    except:
        return text

async def grade_batch_with_gemini(batch_data: List[Dict]) -> Dict[str, Dict[str, Any]]:
    """
    Grades a batch using Gemini with Smart Rubric Splitting and Feedback.
    - Fixes "Student 3 Error" by explicitly splitting Option A/B in the prompt.
    - Returns: { "11": {"score": 3.5, "feedback": "..."} }
    """
    if not batch_data or not gemini_model: 
        logger.error("Skipping AI Grading: No Data or No Model.")
        return {}

    prompt = """Act as a strict academic evaluator. 
    I will provide a list of questions, the correct rubric, and the student's answer.
    
    GLOBAL GRADING RULES:
    1. Use the mark breakdowns in the rubric to assign points (increments of 0.5).
    2. STRICT LABEL MATCHING is required for "Either/Or" questions.
    
    SCENARIO A: RUBRIC HAS "OPTION A" AND "OPTION B" (Explicit Split)
       - If Student Label = 'a'/'A' -> Grade ONLY against OPTION A. (If content matches B, Score 0).
       - If Student Label = 'b'/'B' -> Grade ONLY against OPTION B. (If content matches A, Score 0).
       - If Student Label = Neutral (e.g. '16', 'Ans') or No Label -> Grade against BOTH, pick HIGHER score.
       
    SCENARIO B: RUBRIC IS A SINGLE BLOCK (No Explicit Split)
       - The rubric might list options implicitly (e.g. "1. Definition... 2. Explanation...").
       - Identify which part corresponds to 'a' and 'b'.
       - Verify if the Student's Label matches their Answer Content.
       - CRITICAL: If you are unsure about the boundary or the label is neutral, use CHAMPION LOGIC (Grade against the most relevant part and assign the HIGHER score).
    
    FEEDBACK REQUIREMENTS:
    - Provide brief, constructive feedback (max 2 sentences).
    - Mention exactly why marks were lost (e.g., "Missed key formula").
    
    ITEMS TO GRADE:
    """
    
    for item in batch_data:
        raw_rubric = item['rubric']
        
        # --- PYTHON RUBRIC SPLITTING (The Fix) ---
        # We split the rubric here so the AI definitely sees two separate blocks
        if "[OR Rubric]" in raw_rubric:
            # Handle variations in the separator
            parts = raw_rubric.split("[OR Rubric]")
            if len(parts) < 2:
                 parts = raw_rubric.split("\n[OR Rubric]: ")
            
            # Clean up the parts
            opt_a = parts[0].replace(":", "").strip()
            opt_b = parts[1].replace(":", "").strip() if len(parts) > 1 else "Content missing"
            
            formatted_rubric = f"--- SCENARIO A (SPLIT) ---\nOPTION A:\n{opt_a}\n\nOPTION B:\n{opt_b}"
        else:
            # Fallback for standard questions or implicit rubrics
            formatted_rubric = f"--- SCENARIO B (MERGED) ---\n{raw_rubric}"

        prompt += f"""
        ---
        [ID: {item['id']}]
        Rubric: 
        {formatted_rubric}
        
        Student Answer: {item['student_ans']}
        Max Marks: {item['max']}
        ---
        """
        
    prompt += """
    OUTPUT: Return strictly a valid JSON object. 
    Keys are IDs. Values are OBJECTS with 'score', 'feedback', and 'topic'.
    Example: {
        "11": {"score": 3.5, "feedback": "Correct definition but missed keywords.", "topic": "Neural Networks"}, 
        "12": {"score": 0.0, "feedback": "Label '12a' mismatch.", "topic": "Optimization"}
    }
    Do NOT use Markdown. Just the JSON string.
    """

    # Use the safe call logic
    response = await call_gemini_api_safe(prompt)
    
    if response:
        try:
            clean_text = response.text.replace("```json", "").replace("```", "").strip()
            clean_text = clean_json_string(clean_text)
            
            # Parse the JSON
            data = json.loads(clean_text)
            
            final_results = {}
            for k, v in data.items():
                try:
                    # Handle response whether it's an object (new) or float (old/fallback)
                    if isinstance(v, dict):
                        raw_score = float(v.get("score", 0.0))
                        feedback = v.get("feedback", "No feedback provided.")
                        topic = v.get("topic", "Unknown Topic")
                    else:
                        raw_score = float(v)
                        feedback = "No feedback provided."
                        topic = "Unknown Topic"
                    
                    final_results[k] = {
                        "score": round(raw_score * 2) / 2,  # Round to nearest 0.5
                        "feedback": feedback,
                        "topic": topic
                    }
                except:
                    final_results[k] = {"score": 0.0, "feedback": "Error parsing AI response", "topic": "Unknown Topic"}
            return final_results
        except Exception as e:
            logger.error(f"JSON Parse Error: {e}")
            return {}
    return {}

def fallback_grade_with_minilm(batch_data: List[Dict]) -> Dict[str, float]:
    if not minilm_model: return {}
    scores = {}
    for item in batch_data:
        try:
            emb1 = minilm_model.encode(item['rubric'], convert_to_tensor=True)
            emb2 = minilm_model.encode(item['student_ans'], convert_to_tensor=True)
            sim = util.pytorch_cos_sim(emb1, emb2).item()
            final_score = sim * item['max']
            scores[item['id']] = round(final_score, 2)
        except:
            scores[item['id']] = 0.0
    return scores

# --- Routes ---

@evaluator_app.get("/get-metadata")
async def get_metadata():
    """
    Fetches Departments AND Details so the frontend can populate 
    dependent dropdowns (Batch, Semester, Subject).
    """
    try:
        return await db_module.get_metadata()
    except Exception as e:
        logger.error(f"Metadata Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@evaluator_app.post("/get-students")
async def get_students_endpoint(request: StudentListRequest):
    try:
        dept = request.department.strip()
        batch = request.batch.strip()
        
        logger.info(f"Looking for Students -> Dept: '{dept}', Batch: '{batch}'")

        students = await db_module.get_students(dept, batch)

        if students:
            logger.info(f"✅ Found {len(students)} students.")
            return {"students": students, "exam_id": request.exam_id}
        else:
            logger.warning("❌ No student record found in DB.")
            return {"students": [], "exam_id": request.exam_id}

    except Exception as e:
        logger.error(f"Database Error in get_students: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@evaluator_app.post("/upload-files")
async def upload_files(
    exam_id: str = Form(...),
    question_paper: UploadFile = File(...),
    answer_key: UploadFile = File(...),
    student_papers: List[UploadFile] = File(...)
):
    upload_dir = Path(__file__).parent / "uploads" / exam_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    files = {}
    
    qp_path = upload_dir / question_paper.filename
    with open(qp_path, "wb") as b: shutil.copyfileobj(question_paper.file, b)
    files["question_paper"] = str(qp_path)

    ak_path = upload_dir / answer_key.filename
    with open(ak_path, "wb") as b: shutil.copyfileobj(answer_key.file, b)
    files["answer_key"] = str(ak_path)

    sp_paths = []
    for paper in student_papers:
        p_path = upload_dir / paper.filename
        with open(p_path, "wb") as b: shutil.copyfileobj(paper.file, b)
        sp_paths.append(str(p_path))
    files["student_papers"] = sp_paths

    return {"message": "Files uploaded successfully", "files": files}

async def extract_topics_metadata(qp_map: Dict[str, Dict]) -> Dict[str, str]:
    """
    Extracts topics for ALL questions in the QP map using a single AI call.
    Returns: {"1": "Calculus", "2": "Algebra", ...}
    """
    if not qp_map or not gemini_model: return {}
    
    # Construct Prompt
    q_list_str = ""
    for qid, data in qp_map.items():
        text = data.get('text', '')[:200] # Truncate likely sufficient
        q_list_str += f"Q{qid}: {text}\n"
        
    prompt = f"""
    Act as an academic classifier. 
    I will provide a list of questions from an exam paper.
    Identify the core academic TOPIC for each question (e.g. "Neural Networks", "Optimization", "Calculus").
    Keep topics concise (1-3 words).
    
    QUESTIONS:
    {q_list_str}
    
    OUTPUT:
    Return strictly a valid JSON object. Keys are QIDs (e.g. "1", "11"). Values are the Topic strings.
    Example: {{"1": "Calculus", "11": "Neural Networks"}}
    """
    
    try:
        response = await call_gemini_api_safe(prompt)
        if response:
            clean = clean_json_string(response.text.replace("```json", "").replace("```", "").strip())
            return json.loads(clean)
    except Exception as e:
        logger.error(f"Topic Extraction Failed: {e}")
    
    return {}

@evaluator_app.post("/evaluate")
async def evaluate(
    exam_id: str = Form(...),
    question_paper_path: str = Form(...),
    answer_key_path: str = Form(...),
    student_papers_paths_str: str = Form(..., alias="student_papers_paths"),
    exam_type: str = Form("CIA"),
    subject: str = Form(None),
    batch: str = Form(None),
    department: str = Form(None),
    semester: str = Form(None),
):
    async def evaluation_stream():
        try:
            student_paths = json.loads(student_papers_paths_str)
            logger.info(f"[EVALUATE] Starting evaluation for exam_id={exam_id}, {len(student_paths)} students")
            
            qp_map = parse_docx_table_data(question_paper_path, True)
            key_map = parse_docx_table_data(answer_key_path, False)
            
            # --- TOPIC EXTRACTION (Centralized) ---
            logger.info("Extracting Topics from Question Paper...")
            topic_metadata = await extract_topics_metadata(qp_map)
            logger.info(f"Extracted Topics: {topic_metadata}")
            
            selected_type = "Model" if exam_type == "Models" else exam_type
            schema = EXAM_PATTERNS.get(selected_type, EXAM_PATTERNS["CIA"])
            
            # Ranges
            mcq_ids = [str(i) for i in range(schema["mcq"]["start"], schema["mcq"]["start"] + schema["mcq"]["count"])]
            short_ids = [str(i) for i in range(schema["short"]["start"], schema["short"]["start"] + schema["short"]["count"])]
            long_ids = [str(i) for i in range(schema["long"]["start"], schema["long"]["start"] + schema["long"]["count"])]

            results = []
            total_students = len(student_paths)
            
            for idx, s_path in enumerate(student_paths):
                # --- NEW: VISION GRADING FOR PDF ---
                if s_path.lower().endswith(".pdf"):
                    roll_no = f"Student_{idx+1}" # Fallback
                    ocr_text = ""
                    
                    try:
                        # 1. Attempt extracting from filename
                        fname = os.path.basename(s_path)
                        rn_match = re.search(r"(\d{5,})", fname)
                        if rn_match: 
                            roll_no = rn_match.group(1)
                        else:
                            # 2. Fallback: OCR First Page for Handwritten Identity
                            logger.info(f"Filename extraction failed. Running OCR on first page of {fname}...")
                            ocr_text = extract_first_page_text_ocr(s_path)
                            extracted_id = extract_student_identity(ocr_text)
                            if extracted_id:
                                roll_no = extracted_id
                                logger.info(f"✅ Extracted Roll No via OCR: {roll_no}")
                            else:
                                logger.warning(f"❌ Could not identify student in {fname}")
                    except Exception as e:
                        logger.error(f"Identity Extraction Error: {e}")

                    logger.info(f"[EVALUATE] ⚡ Vision Grading for {roll_no} (PDF)...")
                    
                    # 1. Build Full Rubric String
                    full_rubric_str = "--- MASTER RUBRIC ---\n"
                    
                    # Add MCQs (Optional: If we want AI to grade visual MCQs too)
                    # For now, let's include everything
                    all_ids = mcq_ids + short_ids + long_ids
                    
                    for q_id in all_ids:
                        if q_id in key_map:
                            q_text = qp_map.get(q_id, {}).get("text", "Question Text Missing")
                            r_text = key_map[q_id]["text"]
                            max_m = 1.0
                            if q_id in short_ids: max_m = schema["short"]["marks"]
                            elif q_id in long_ids: max_m = schema["long"]["marks"]
                            
                            full_rubric_str += f"\n[Q{q_id}] (Max: {max_m})\nQuestion: {q_text}\nRubric: {r_text}\n"

                    # 2. Call Vision API
                    # Use Gemini 2.0 Flash or 1.5 Flash (User asked for 3, but let's stick to stable/available)
                    # We can try to respect user wish: 'gemini-2.0-flash-exp' or 'gemini-1.5-flash'
                    # The library usually handles model aliases.
                    vision_results = await grade_pdf_with_vision(s_path, full_rubric_str, model_name="gemini-3-flash-preview")
                    
                    # 3. Process Results
                    marks = {}
                    feedback = {}
                    topics = {}
                    total_score = 0
                    
                    if vision_results:
                        for q_id, res in vision_results.items():
                            # clean key "11" -> "Q11"
                            clean_id = q_id.replace("Q", "")
                            key = f"Q{clean_id}"
                            
                            score = float(res.get("score", 0.0))
                            fb = res.get("feedback", "")
                            
                            marks[key] = score
                            feedback[key] = fb
                            
                            # Prioritize specific topic from Vision AI (which sees the specific answer choice)
                            # Fallback to centralized topic metadata
                            specific_topic = res.get("topic", "Unknown")
                            central_topic = topic_metadata.get(clean_id, "General")
                            
                            if specific_topic != "Unknown":
                                topics[key] = specific_topic
                            else:
                                topics[key] = central_topic
                            
                            total_score += score
                    else:
                        logger.error(f"Vision grading returned empty for {roll_no}")
                        feedback["General"] = "Vision Grading Failed. Please check logs."

                # --- OLD: TEXT GRADING FOR DOCX ---
                else: 
                    s_text = extract_text(s_path)
                    
                    # 1. Try Filename First
                    fname = os.path.basename(s_path)
                    rn_match = re.search(r"(\d{5,})", fname)
                    
                    if rn_match:
                        roll_no = rn_match.group(1)
                    else:
                        # 2. Fallback to Content
                        roll_no = extract_student_identity(s_text) or f"UNKNOWN_{idx}"
                        
                    s_answers = parse_student_text(s_text)
                    
                    marks = {}
                    feedback = {}  # Store feedback per question
                    topics = {}
                    total_score = 0
                    master_batch = []
    
                    # --- 1. LOCAL GRADING: MCQs (0 API COST) ---
                    for q_id in mcq_ids:
                        if q_id in key_map:
                            # Clean extraction to prevent whitespace errors
                            raw_key = key_map[q_id]['text'].lstrip("- ").strip()
                            raw_student = s_answers.get(q_id, "").lstrip("- ").strip()
                            
                            model_char = raw_key[0].upper() if raw_key else "X"
                            student_char = raw_student[0].upper() if raw_student else "Y"
                            
                            if model_char == student_char:
                                score = 1.0
                                feedback[f"Q{q_id}"] = "Correct"
                            else:
                                score = 0.0
                                feedback[f"Q{q_id}"] = f"Incorrect. Correct answer: {model_char}"
                            
                            marks[f"Q{q_id}"] = score
                            topics[f"Q{q_id}"] = topic_metadata.get(q_id, "General") # Use Centralized Metadata
                            total_score += score
    
                    # --- 2. AI PREPARATION: All Descriptive Questions (Master Batch) ---
                    all_descriptive_ids = short_ids + long_ids
                    
                    for q_id in all_descriptive_ids:
                        if q_id in key_map:
                            max_m = schema["short"]["marks"] if q_id in short_ids else schema["long"]["marks"]
                            
                            # Add to the single master list
                            master_batch.append({
                                "id": q_id,
                                "question": qp_map.get(q_id, {}).get("text", ""),
                                "rubric": key_map[q_id]["text"], 
                                "student_ans": s_answers.get(q_id, "No Answer"),
                                "max": max_m
                            })
    
                    # --- 3. SINGLE API CALL PER STUDENT (Scores + Feedback) ---
                    if master_batch:
                        logger.info(f"[EVALUATE] 🚀 Master Call for {roll_no}: Grading {len(master_batch)} questions")
                        
                        # Call Gemini ONCE - returns {"qid": {"score": X, "feedback": "..."}}
                        ai_results = await grade_batch_with_gemini(master_batch)
                        
                        # Distribute scores and feedback
                        for item in master_batch:
                            qid = item['id']
                            result_data = ai_results.get(qid, {"score": 0.0, "feedback": "", "topic": "Unknown"})
                            
                            final_val = float(result_data.get("score", 0.0))
                            marks[f"Q{qid}"] = final_val
                            feedback[f"Q{qid}"] = result_data.get("feedback", "")
                            
                            # Use Centralized Metadata + Fallback to AI result
                            # EDIT: Prioritize AI result (specific to student answer for choice questions)
                            central_topic = topic_metadata.get(qid)
                            ai_topic = result_data.get("topic", "Unknown")
                            
                            if ai_topic != "Unknown":
                                topics[f"Q{qid}"] = ai_topic
                            else:
                                topics[f"Q{qid}"] = central_topic if central_topic else "General"
                            
                            total_score += final_val
                
                # --- Finish Student (Include Feedback & Metadata) ---
                res = {
                    "roll_no": roll_no, 
                    "exam_id": exam_id, 
                    "marks": marks, 
                    "feedback": feedback,
                    "topics": topics,
                    "total": round(total_score, 2), 
                    "timestamp": datetime.datetime.utcnow().isoformat(),
                    # Store Metadata
                    "subject": subject,
                    "batch": batch,
                    "department": department,
                    "semester": semester,
                    "exam_type": exam_type
                }
                await db_module.insert_evaluation(res)
                res["_id"] = str(res.get("_id", ""))
                results.append(res)
                yield json.dumps({"type": "progress", "value": int(((idx + 1) / total_students) * 100), "message": f"Graded {roll_no}"}) + "\n"

            yield json.dumps({"type": "complete", "status": "success", "results": results}) + "\n"
        except Exception as e:
            logger.error(f"Evaluation Error: {e}")
            yield json.dumps({"type": "error", "message": str(e)}) + "\n"
    
    return StreamingResponse(evaluation_stream(), media_type="application/x-ndjson")

@evaluator_app.post("/update-marks")
async def update_marks(request: MarkUpdateRequest):
    try:
        evaluation = await db_module.find_evaluation(request.exam_id, request.roll_no)
        if not evaluation: raise HTTPException(404, detail="Not found")
        q_key = request.question_num if request.question_num.startswith("Q") else f"Q{request.question_num}"
        evaluation["marks"][q_key] = request.new_mark
        new_total = sum(evaluation["marks"].values())
        await db_module.update_evaluation_marks(int(evaluation["_id"]), evaluation["marks"], new_total)
        return {"status": "success"}
    except Exception as e:
        raise HTTPException(500, detail=str(e))

@evaluator_app.delete("/delete-evaluation/{exam_id}")
async def delete_evaluation_endpoint(exam_id: str):
    try:
        await db_module.delete_evaluation(exam_id)
        return {"status": "success", "message": "Evaluation deleted"}
    except Exception as e:
        logger.error(f"Error deleting evaluation {exam_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@evaluator_app.get("/history")
async def get_evaluation_history():
    try:
        return await db_module.get_evaluation_history()
    except Exception as e:
        logger.error(f"Error fetching evaluation history: {e}")
        return []

@evaluator_app.get("/results/{exam_id}")
async def get_evaluation_results(exam_id: str):
    try:
        return await db_module.get_evaluation_results(exam_id)
    except Exception as e:
        logger.error(f"Error fetching results for {exam_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@evaluator_app.get("/export-excel")
async def export_excel(exam_id: str):
    try:
        evals_raw = await db_module.get_evaluation_results(exam_id)
        if not evals_raw: raise HTTPException(404, detail="No data found")
        
        # Determine Exam Pattern (CIA vs Model) from the first record metadata if available, 
        # or guess based on question count/keys.
        # Ideally, we should store 'exam_type' in evaluations table, but it's not strictly there.
        # However, we can infer Max Score from Question Number.
        
        # LOGIC:
        # CIA: Q1-10 (1), Q11-15 (4), Q16-17 (10)
        # Model: Q1-10 (1), Q11-15 (5), Q16-20 (8)
        
        # Let's try to detect if it's CIA or Model based on Q11-15 max scores? 
        # Actually, the user rules are explicit. But we don't know the Exam Type for sure here unless we look at the range of questions present.
        # If Q18 exist -> Likely Model (since CIA stops at 17).
        # Let's check keys of the first student.
        
        first_keys = evals_raw[0]["marks"].keys()
        max_q_num = 0
        for k in first_keys:
            num = int(re.search(r'\d+', k).group())
            if num > max_q_num: max_q_num = num
            
        is_model = max_q_num > 17
        exam_mode = "Model" if is_model else "CIA"
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Detailed Report"
        
        # Headers: Roll No | Question | Topic | Score | Max_Score | Feedback
        headers = ["Roll No", "Question", "Topic", "Score", "Max_Score", "Feedback"]
        header_fill = PatternFill(start_color="3B82F6", end_color="3B82F6", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.fill = header_fill
            cell.font = header_font
            
        row_idx = 2
        
        for record in evals_raw:
            roll_no = record.get("roll_no", "Unknown")
            marks_map = record.get("marks", {})
            feedback_map = record.get("feedback", {})
            topics_map = record.get("topics", {})
            
            # Sort questions naturally (Q1, Q2... Q10)
            sorted_qs = sorted(marks_map.keys(), key=lambda x: int(re.search(r'\d+', x).group()) if re.search(r'\d+', x) else 999)
            
            for q_key in sorted_qs:
                q_num = int(re.search(r'\d+', q_key).group())
                
                # Determine Max Score
                max_score = 0
                if 1 <= q_num <= 10:
                    max_score = 1
                elif 11 <= q_num <= 15:
                    max_score = 5 if exam_mode == "Model" else 4
                elif q_num >= 16:
                    max_score = 8 if exam_mode == "Model" else 10
                    
                score = marks_map.get(q_key, 0)
                topic = topics_map.get(q_key, "General")
                fb = feedback_map.get(q_key, "")
                
                ws.cell(row=row_idx, column=1, value=roll_no)
                ws.cell(row=row_idx, column=2, value=q_key)
                ws.cell(row=row_idx, column=3, value=topic)
                ws.cell(row=row_idx, column=4, value=score)
                ws.cell(row=row_idx, column=5, value=max_score)
                ws.cell(row=row_idx, column=6, value=fb)
                
                row_idx += 1
                
        # Auto-width
        dims = {1: 15, 2: 10, 3: 25, 4: 8, 5: 10, 6: 40}
        for col, width in dims.items():
            ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = width
            
        path = f"results_{exam_id}_detailed.xlsx"
        wb.save(path)
        return FileResponse(path, filename=path)

    except Exception as e:
        raise HTTPException(500, detail=str(e))

# --- STATIC FILES (Self-contained) ---
if os.path.exists("evaluator"):
    evaluator_app.mount("/", StaticFiles(directory="evaluator", html=True), name="evaluator_static")
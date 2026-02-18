import os
import re
import json
import random
import asyncio
import datetime
from pathlib import Path
from dotenv import load_dotenv
import io
import logging
import time

# Load environment variables from backend/.env
load_dotenv(Path(__file__).parent.parent / ".env")

from fastapi import FastAPI, Request, HTTPException, UploadFile, File, Form
from fastapi.responses import RedirectResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Dict, Any

import core.database as db_module
import services.question_generator as qg
import routers.learning as learning_routes

import bcrypt
from pydantic import BaseModel
from docx import Document
from docx.text.paragraph import Paragraph

# --- IMPORT EVALUATOR ---
from routers.evaluator import evaluator_app

# --- CONFIGURATION ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

BASE_DIR = Path(__file__).resolve().parent
TEMP_DIR = BASE_DIR / "temp"
TEMP_DIR.mkdir(exist_ok=True)

TEMPLATE_PATHS = {
    "CIA": {
        "paper": BASE_DIR / "templates" / "CIA_QP_template.docx",
        "key": BASE_DIR / "templates" / "CIA_Answerkey_template.docx"
    },
    "Model": {
        "paper": BASE_DIR / "templates" / "Models_QP_template.docx",
        "key": BASE_DIR / "templates" / "Models_Answerkey_template.docx"
    }
}

# --- FASTAPI APP SETUP ---
app = FastAPI(title="Question Paper Generator", version="2.0")

# --- CORS ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- ROUTERS ---
app.include_router(learning_routes.router, prefix="/learning")

# --- Database lifecycle (PostgreSQL via asyncpg) ---
# --- AUTH CONFIG ---
# JWT Removed as per request

# --- ROUTES ---
# --- PYDANTIC MODELS ---
class Question(BaseModel):
    id: int
    type: str
    text: str
    answer: str
    marks: int

class RegenerateRequest(BaseModel):
    current_question: Question
    subject: str
    difficulty: str
    topics: List[str] = []

class DownloadRequest(BaseModel):
    department: str
    batch: str
    semester: str
    subject: str
    examType: str
    duration: str
    paperSetter: str
    hod: str
    questions: List[Question]

@app.get("/")
async def read_root(): return {"message": "KCLAS Question Paper Generator API", "version": "2.1"}

@app.post("/login")
async def login(request: Request, email: str = Form(...), password: str = Form(...)):
    """
    Login endpoint. Returns user info directly (No JWT).
    """
    try:
        user = await db_module.find_user_by_email(email)
        if user and bcrypt.checkpw(password.encode('utf-8'), user["password"].encode('utf-8')):
            role = user.get("role", "faculty")
            
            return {
                "message": "Login successful",
                "role": role,
                "email": email,
                "user_id": user["id"],
                "redirect": "/dashboard" if role == "faculty" else "/student-dashboard"
            }
        
        raise HTTPException(status_code=401, detail="Invalid credentials")
    except HTTPException: raise
    except Exception as e:
        logger.error(f"Login error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail="Login failed")

@app.get("/departments")
async def get_departments(): return await db_module.get_departments()

@app.get("/details/{department}")
async def get_details_by_department(department: str): return await db_module.get_details_by_department(department.strip())

@app.get("/saved-papers")
async def get_saved_papers(): return await db_module.get_saved_papers()

@app.get("/saved-paper/{paper_id}")
async def get_saved_paper(paper_id: str):
    paper = await db_module.get_saved_paper(int(paper_id))
    if not paper: raise HTTPException(404, "Paper not found")
    return paper

@app.post("/upload-syllabus")
async def upload_syllabus(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    temp_file_path = TEMP_DIR / f"temp_{random.randint(1000, 9999)}_{file.filename}"
    try:
        with open(temp_file_path, "wb") as f: f.write(await file.read())
        # Use QG
        units = await qg.extract_units_from_pdf(str(temp_file_path))
        return {"units": units}
    finally:
        if temp_file_path.exists(): temp_file_path.unlink()

@app.post("/generate-questions")
async def generate_questions(request: Request):
    try:
        data = await request.json()
        exam_type = "Model" if data.get("exam_type") == "Models" else data.get("exam_type")
        
        # Use QG
        doc_chunks = qg.create_document_chunks(data.get("selected_units", []))
        if not doc_chunks: raise HTTPException(status_code=422, detail="Failed to process content.")
        
        # Get selected topics
        topics = data.get("selected_topics", []) # List[str]

        paper = await qg.generate_question_paper(
            tuple(doc.page_content for doc in doc_chunks), 
            data["subject"], 
            exam_type, 
            data["difficulty"],
            topics 
        )
        
        await db_module.insert_question_paper(
            subject=data["subject"],
            exam_type=exam_type,
            difficulty=data["difficulty"],
            paper=paper
        )
        return {"question_paper": paper}
    except Exception as e:
        logger.error(f"Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/regenerate-question")
async def regenerate_question(req: RegenerateRequest):
    try:
        # Create a specific prompt for regeneration
        q_type = req.current_question.type
        marks = req.current_question.marks
        
        prompt = qg.RUBRIC_BATCH_PROMPT if q_type != "MCQ" else qg.MCQ_BATCH_PROMPT
        q_label = "MCQ" if q_type == "MCQ" else ("Short Answer" if q_type == "Short Answer" else "Long Essay")
        
        # Reuse run_batch_query but ask for 1 question
        context_override = f"Original Question: {req.current_question.text}\nTask: Generate a SIMILAR but DISTINCT variation of this question."
        
        new_questions = await qg.run_batch_query(
            prompt, q_label, 1, marks, context_override, req.subject, req.difficulty
        )
        
        if not new_questions: raise ValueError("Failed to regenerate")
        
        new_q = new_questions[0]
        # Map back to Question model
        return Question(
            id=req.current_question.id,
            type=q_type,
            text=new_q.get("text", "Error"),
            answer=new_q.get("answer", "Error"),
            marks=marks
        )
    except Exception as e:
        logger.error(f"Regeneration failed: {e}")
        raise HTTPException(status_code=500, detail="Regeneration failed")

# --- DOCX UTILS ---
import latex2mathml.converter
from lxml import etree

try:
    # Load transformation stylesheet
    xslt_tree = etree.parse(str(BASE_DIR / "resources" / "MML2OMML.xsl"))
    xslt_transform = etree.XSLT(xslt_tree)
except Exception as e:
    logger.error(f"Failed to load MML2OMML.xsl: {e}")
    xslt_transform = None

def latex_to_omml(latex_str):
    if not xslt_transform: return None
    try:
        mathml = latex2mathml.converter.convert(latex_str)
        mathml_tree = etree.fromstring(mathml)
        omml_tree = xslt_transform(mathml_tree)
        return omml_tree
    except Exception as e:
        logger.error(f"Math conversion failed for '{latex_str}': {e}")
        return None

def replace_text_in_paragraph(paragraph: Paragraph, context: Dict[str, str]):
    full_text = "".join(run.text for run in paragraph.runs)
    if '{{' not in full_text: return
    
    # Perform Replacement
    for key, value in context.items():
        if key in full_text: full_text = full_text.replace(key, str(value))
    
    # Check if there is any LaTeX to render ($...$)
    segments = re.split(r'(\$.*?\$)', full_text)
    
    # Clear existing runs
    p = paragraph._p
    for run in paragraph.runs:
        p.remove(run._r)
        
    for segment in segments:
        if not segment: continue
        
        # Check if math
        if segment.startswith('$') and segment.endswith('$') and len(segment) > 2:
            latex_content = segment[1:-1] # Strip $
            omml_element = latex_to_omml(latex_content)
            
            if omml_element is not None:
                from docx.oxml import parse_xml
                omml_xml_str = etree.tostring(omml_element, encoding='unicode')
                try:
                    oxml_obj = parse_xml(omml_xml_str)
                    paragraph._p.append(oxml_obj)
                except Exception as ex:
                    logger.error(f"Failed to insert OMML: {ex}")
                    paragraph.add_run(segment) # Fallback
            else:
                 paragraph.add_run(segment) # Fallback text
        else:
            # Regular text
            if '\n' in segment:
                lines = segment.split('\n')
                for i, line in enumerate(lines):
                    if i > 0: paragraph.add_run().add_break()
                    paragraph.add_run(line)
            else:
                paragraph.add_run(segment)

def replace_placeholders(doc: Document, context: Dict[str, str]):
    for p in doc.paragraphs: replace_text_in_paragraph(p, context)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: replace_text_in_paragraph(p, context)
    for section in doc.sections:
        for p in section.header.paragraphs: replace_text_in_paragraph(p, context)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: replace_text_in_paragraph(p, context)
        for p in section.footer.paragraphs: replace_text_in_paragraph(p, context)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: replace_text_in_paragraph(p, context)
    return doc

def parse_answer_and_marks(answer_text: str):
    if not answer_text: return "", ""
    lines = answer_text.strip().split('\n')
    cleaned_answer_lines = []
    formatted_marks_lines = []
    mark_pattern = re.compile(r'\s*\((\d+(?:\.\d+)?)(?:\s*marks?)?\)\s*$', re.IGNORECASE)

    for line in lines:
        line = line.strip()
        if line.startswith('**Keywords') or line.startswith('Keywords'):
            cleaned_answer_lines.append(line)
            formatted_marks_lines.append("")
            continue
            
        match = mark_pattern.search(line)
        if match:
            mark_number = match.group(1)
            formatted_marks_lines.append(f"({mark_number})")
            cleaned_line = mark_pattern.sub('', line).strip()
            cleaned_answer_lines.append(cleaned_line)
        else:
            cleaned_answer_lines.append(line)
            formatted_marks_lines.append("")

    return "\n".join(cleaned_answer_lines), "\n".join(formatted_marks_lines)


@app.post("/download-paper")
async def download_paper(data: DownloadRequest):
    exam_type = "Model" if data.examType == "Models" else data.examType
    template_path = TEMPLATE_PATHS.get(exam_type, {}).get("paper")
    if not template_path or not template_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")

    doc = Document(template_path)
    context = {
        "{{department}}": data.department, "{{batch}}": data.batch, "{{semester}}": data.semester,
        "{{subject}}": data.subject, "{{examType}}": data.examType, "{{duration}}": data.duration,
        "{{paperSetter}}": data.paperSetter, "{{hod}}": data.hod,
    }
    mcqs = [q for q in data.questions if q.type == "MCQ"]
    short_answers = [q for q in data.questions if q.type == "Short Answer"]
    long_essays = [q for q in data.questions if q.type == "Long Essay"]

    for i in range(1, 11):
        if i <= len(mcqs):
            q = mcqs[i - 1]
            parts = q.text.split('\n')
            context[f"{{{{Q{i}}}}}"] = parts[0]
            context[f"{{{{Q{i}_A}}}}"] = parts[1] if len(parts) > 1 else ""
            context[f"{{{{Q{i}_B}}}}"] = parts[2] if len(parts) > 2 else ""
            context[f"{{{{Q{i}_C}}}}"] = parts[3] if len(parts) > 3 else ""
            context[f"{{{{Q{i}_D}}}}"] = parts[4] if len(parts) > 4 else ""
    
    q_num = 11
    for i in range(0, len(short_answers), 2):
        context[f"{{{{Q{q_num}a}}}}"] = short_answers[i].text
        context[f"{{{{Q{q_num}b}}}}"] = short_answers[i + 1].text if i + 1 < len(short_answers) else ""
        q_num += 1

    q_num = 16
    for i in range(0, len(long_essays), 2):
        context[f"{{{{Q{q_num}a}}}}"] = long_essays[i].text
        context[f"{{{{Q{q_num}b}}}}"] = long_essays[i + 1].text if i + 1 < len(long_essays) else ""
        q_num += 1

    doc = replace_placeholders(doc, context)
    file_path = TEMP_DIR / f"{data.subject}_QP_{random.randint(1000,9999)}.docx"
    doc.save(file_path)
    return FileResponse(path=file_path, filename=f"{data.subject}_Question_Paper.docx")

@app.post("/download-key")
async def download_key(data: DownloadRequest):
    exam_type = "Model" if data.examType == "Models" else data.examType
    template_path = TEMPLATE_PATHS.get(exam_type, {}).get("key")
    if not template_path or not template_path.exists():
        raise HTTPException(status_code=404, detail="Key template not found")
    
    doc = Document(template_path)
    context = {
        "{{department}}": data.department, "{{batch}}": data.batch, "{{semester}}": data.semester,
        "{{subject}}": data.subject, "{{examType}}": data.examType, "{{duration}}": data.duration,
        "{{paperSetter}}": data.paperSetter, "{{hod}}": data.hod,
    }
    mcqs = [q for q in data.questions if q.type == "MCQ"]
    short = [q for q in data.questions if q.type == "Short Answer"]
    long = [q for q in data.questions if q.type == "Long Essay"]

    for i in range(1, 11):
        context[f"{{{{A{i}}}}}"] = mcqs[i-1].answer if i <= len(mcqs) else ""

    q_num = 11
    for i in range(0, len(short), 2):
        a, m = parse_answer_and_marks(short[i].answer.replace('\\n', '\n'))
        context[f"{{{{A{q_num}a}}}}"] = a
        context[f"{{{{M{q_num}a}}}}"] = m
        if i + 1 < len(short):
            a2, m2 = parse_answer_and_marks(short[i+1].answer.replace('\\n', '\n'))
            context[f"{{{{A{q_num}b}}}}"] = a2
            context[f"{{{{M{q_num}b}}}}"] = m2
        q_num += 1

    q_num = 16
    for i in range(0, len(long), 2):
        a, m = parse_answer_and_marks(long[i].answer.replace('\\n', '\n'))
        context[f"{{{{A{q_num}a}}}}"] = a
        context[f"{{{{M{q_num}a}}}}"] = m
        if i + 1 < len(long):
            a2, m2 = parse_answer_and_marks(long[i+1].answer.replace('\\n', '\n'))
            context[f"{{{{A{q_num}b}}}}"] = a2
            context[f"{{{{M{q_num}b}}}}"] = m2
        q_num += 1
        
    doc = replace_placeholders(doc, context)
    file_path = TEMP_DIR / f"{data.subject}_Key_{random.randint(1000,9999)}.docx"
    doc.save(file_path)
    return FileResponse(path=file_path, filename=f"{data.subject}_Answer_Key.docx")

if __name__ == "__main__":
    import uvicorn
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--port", type=int, default=8000)
    args = parser.parse_args()
    uvicorn.run(app, host="0.0.0.0", port=args.port)
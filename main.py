import os
import re
import json
import random
import asyncio
import datetime
from pathlib import Path
from dotenv import load_dotenv
import io
import logging
import time

# Load environment variables
load_dotenv()

from fastapi import FastAPI, Request, HTTPException, UploadFile, File, Form
from fastapi.responses import RedirectResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Dict, Any

import database as db_module

import pdfplumber
import bcrypt
from pydantic import BaseModel
from docx import Document
from docx.text.paragraph import Paragraph

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI

# --- IMPORT EVALUATOR ---
# This pulls the 'evaluator_app' FastAPI object from main2.py
from main2 import evaluator_app

# --- CONFIGURATION ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- AI & DB INITIALIZATION ---
try:
    GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
    if not GOOGLE_API_KEY:
        raise ValueError("GOOGLE_API_KEY environment variable not set")
    os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY
    llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.6)
    embedding = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
    logger.info("Google AI Models initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize Google AI Models. Error: {e}")
    raise


# --- PROMPT ENGINEERING ---
MCQ_BATCH_PROMPT = PromptTemplate.from_template("""
You are an expert question paper setter for {subject}. 
Generate **{num_questions}** distinct Multiple Choice Questions of {difficulty} difficulty, each worth {marks} mark.

For each MCQ, provide:
- "question": The question text
- "options": An object with keys A, B, C, D and their option texts
- "answer": The correct option letter (A, B, C, or D)

**IMPORTANT**: For subjects like **Mathematics, Statistics, or Physics**:
- The questions MUST be **numerical problem-solving** or **calculation-based**.
- Avoid purely theoretical definitions or "What is" questions unless absolutely necessary.
- For mathematical symbols, equations, or formulas, use standard LaTeX formatting enclosed in single dollar signs like $E=mc^2$. 

Return a single, valid JSON object:
{{"questions": [
  {{"question": "...", "options": {{"A": "...", "B": "...", "C": "...", "D": "..."}}, "answer": "A"}},
  ...
]}}

**Context:** {context}
""")

RUBRIC_BATCH_PROMPT = PromptTemplate.from_template("""
You are an expert question paper setter for **{subject}**.
Generate **{num_questions}** distinct {question_type} questions of {difficulty} difficulty, each worth {marks} marks.

For each question, provide:
- "question": The question text
- "answer": A detailed marking scheme/rubric in this EXACT format:
  * **For Math/Statistics**: Provide **detailed step-by-step calculations** to solve the problem.
  * **For Theory**: Use multiple bullet points.
  * **CRITICAL**: If the subject is Math or Statistics, generate **NUMERICAL PROBLEMS** that require solving, not just explaining concepts.
  * For mathematical symbols, equations, or formulas, use standard LaTeX formatting enclosed in single dollar signs like $E=mc^2$.
  * Format:
    - A specific assessment criterion (what the student should include)
    - The mark allocation for that point in parentheses, e.g. "(3)"
  * The marks should sum to {marks}
  * End with a "Keywords:" line listing 4-6 important terms

Example answer format for a {marks}-mark question:
"- Provides a clear definition of the concept (3)\\n- Explains key characteristics and differences (3)\\n- Discusses practical applications or examples (4)\\n**Keywords:** Term1, Term2, Term3, Term4"

Return a single, valid JSON object:
{{"questions": [
  {{"question": "...", "answer": "- Point 1 (marks)\\n- Point 2 (marks)\\n**Keywords:** ..."}},
  ...
]}}

**Context:** {context}
""")

# --- FASTAPI APP SETUP ---
app = FastAPI(title="Question Paper Generator", version="2.0")

# --- Database lifecycle (PostgreSQL via asyncpg) ---
@app.on_event("startup")
async def startup():
    await db_module.init_db()
    logger.info("PostgreSQL connection pool ready")

@app.on_event("shutdown")
async def shutdown():
    await db_module.close_db()
    logger.info("PostgreSQL connection pool closed")

# CORS Configuration for Next.js Frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).parent
TEMP_DIR = BASE_DIR / "temp"
TEMP_DIR.mkdir(exist_ok=True)

TEMPLATE_DIR_DOCX = BASE_DIR / "templates"
TEMPLATE_PATHS = {
    "CIA": {
        "paper": TEMPLATE_DIR_DOCX / "CIA_QP_template.docx",
        "key": TEMPLATE_DIR_DOCX / "CIA_Answerkey_template.docx"
    },
    "Model": {
        "paper": TEMPLATE_DIR_DOCX / "Models_QP_template.docx",
        "key": TEMPLATE_DIR_DOCX / "Models_Answerkey_template.docx"
    }
}

# Mount Evaluator (main2.py) at /evaluator
# This allows it to run on the SAME port (8000)
app.mount("/evaluator", evaluator_app)

# --- PYDANTIC MODELS ---
class Question(BaseModel):
    id: int
    type: str
    text: str
    answer: str
    marks: int

class DownloadRequest(BaseModel):
    department: str
    batch: str
    semester: str
    subject: str
    examType: str
    duration: str
    paperSetter: str
    hod: str
    questions: List[Question]

# --- CORE FUNCTIONS ---
import google.generativeai as genai

# Configure GenAI (using same key as LangChain)
if os.getenv("GOOGLE_API_KEY"):
    genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

def extract_units_from_pdf(pdf_path: str) -> List[Dict[str, str]]:
    try:
        full_text = ""
        # 1. Try Standard Text Extraction first
        try:
            with pdfplumber.open(pdf_path) as pdf:
                full_text = "".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as e:
            logger.error(f"PDFPlumber failed: {e}")

        # 2. Check if Scanned (OCR needed)
        if len(full_text) < 100:
            logger.warning("PDF appears to be scanned (text < 100 chars). Attempting Gemini OCR...")
            try:
                # Upload to Gemini
                sample_file = genai.upload_file(path=pdf_path, display_name="Syllabus PDF")
                logger.info(f"Uploaded file to Gemini: {sample_file.uri}")
                
                # Wait for processing
                while sample_file.state.name == "PROCESSING":
                    time.sleep(2)
                    sample_file = genai.get_file(sample_file.name)
                
                if sample_file.state.name == "FAILED":
                    raise ValueError("Gemini file processing failed")

                # Prompt Gemini to extract text
                model = genai.GenerativeModel('gemini-2.5-flash')
                response = model.generate_content([
                    "Extract the full text from this syllabus PDF. Preserve structure like 'Unit 1', 'Module 1'.", 
                    sample_file
                ])
                full_text = response.text
                logger.info("Gemini OCR successful")
                
                # Cleanup
                genai.delete_file(sample_file.name)

            except Exception as e:
                logger.error(f"Gemini OCR Failed: {e}")
                # Fallback to whatever we have (likely empty)
        
        # DEBUG: Save full extracted text to analyze
        debug_path = BASE_DIR / "debug_last_pdf_extraction.txt"
        with open(debug_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        logger.info(f"Saved extracted text to {debug_path}")

        logger.info(f"Extracted {len(full_text)} characters from PDF")
        # Log first 500 chars to see header structure
        logger.info(f"PDF Header Preview: {full_text[:500]}")

        # Improved regex to handle "Unit 1", "Unit I", "Module 1", "Module I"
        # Matches Unit/Module followed by:
        # 1. Digits (\d+)
        # 2. Roman Numerals ([IVX]+)
        unit_pattern = re.compile(
            r"((?:Unit|Module)[\s:-]+(?:(?:\d+)|(?:\b[IVX]+\b)).*?(?=(?:Unit|Module)[\s:-]+(?:(?:\d+)|(?:\b[IVX]+\b))|$))", 
            re.DOTALL | re.IGNORECASE
        )
        
        matches = unit_pattern.findall(full_text)
        
        if not matches:
            logger.info("No explicit units found. Returning full syllabus.")
            return [{"unit": "Full Syllabus", "text": full_text}] if full_text.strip() else []
            
        logger.info(f"Found {len(matches)} units.")
        return [{"unit": f"Unit {idx+1}", "text": match.strip()} for idx, match in enumerate(matches) if match.strip()]
    except Exception as e:
        logger.error(f"Error reading PDF {pdf_path}: {e}")
        return []

def create_document_chunks(units: List[Dict[str, str]]) -> List[Any]:
    texts = [unit["text"] for unit in units if unit.get("text")]
    if not texts: return []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=4000, chunk_overlap=500)
    return text_splitter.create_documents(texts, metadatas=[{"unit": unit["unit"]} for unit in units])

def parse_json_output(response_text: str) -> List[Dict[str, str]]:
    try:
        # Try code block first
        json_match = re.search(r'```(?:json)?\s*(\{[\s\S]*?\})\s*```', response_text, re.DOTALL)
        if not json_match:
            # Try to find JSON object - use greedy match for nested structures
            # Find opening brace and match to the last closing brace
            start = response_text.find('{')
            end = response_text.rfind('}')
            if start != -1 and end != -1 and end > start:
                json_str = response_text[start:end+1]
            else:
                logger.warning(f"No JSON found in response: {response_text[:200]}")
                return []
        else:
            json_str = json_match.group(1)
        
        try:
            data = json.loads(json_str)
        except json.JSONDecodeError as e:
            logger.warning(f"Initial JSON parse failed: {e}. Attempting to fix backslashes...")
            # Fix invalid escapes for LaTeX (e.g., \le -> \\le, \g -> \\g)
            # Strategy: Find ANY backslash.
            # If it's a valid JSON escape (e.g., \n, \", \u1234), keep it.
            # If it's invalid (e.g., \alpha, \underline), escape it (\\alpha, \\underline).
            
            def escape_fixer(match):
                # match.group(0) is the backslash and the following character
                # Check if it looks like a valid escape
                seq = match.group(0) # e.g. "\n" or "\u" or "\a"
                
                # Standard single-char escapes
                if seq in [r'\"', r'\\', r'\/', r'\b', r'\f', r'\n', r'\r', r'\t']:
                    return seq
                    
                # Unicode escape start (\u). We need to peek ahead 4 chars in a real parser,
                # but here we can just regex for \uXXXX in the main pattern.
                # If we matched just "\u", it means it wasn't followed by 4 hex digits (see pattern below).
                return r"\\" + seq[1]

            # Regex:
            # 1. Match valid unicode escape: \\u[0-9a-fA-F]{4} -> Keep as is (Group 1)
            # 2. Match backslash + any char: \\. -> Pass to callback to decide
            # We use re.sub with a function.
            
            # Actually, simpler with re.sub using a smarter pattern:
            # Pattern matches ANY backslash-plus-char sequence.
            # But we prioritize valid unicode first.
            
            def replace_invalid_escapes(text):
                # Pattern 1: Valid Unicode Escape (e.g., \u0020) - Capture it
                # Pattern 2: Any other backslash sequence (e.g., \n, \", \a, \$, \u12 (incomplete)) - Capture it
                pattern = r'(\\u[0-9a-fA-F]{4})|(\\.)'
                
                def sub_callback(m):
                    if m.group(1): return m.group(1) # Valid Unicode, return as is
                    
                    val = m.group(2) # Backslash + char
                    if val in [r'\"', r'\\', r'\/', r'\b', r'\f', r'\n', r'\r', r'\t']:
                        return val # Valid standard escape, keep it
                    
                    # Otherwise (e.g. \$, \a, \u(without hex)), double the backslash
                    return r'\\' + val[1]
                
                return re.sub(pattern, sub_callback, text)

            fixed_json_str = replace_invalid_escapes(json_str)
            try:
                data = json.loads(fixed_json_str)
                logger.info("JSON parse successful after fixing backslashes (Smart Fix).")
            except json.JSONDecodeError as e2:
                logger.error(f"Smart fix failed: {e2}")
                logger.error(f"Snippet: {fixed_json_str[:200]}")
                return []
        
        questions = data.get("questions", [])
        if not isinstance(questions, list):
            logger.warning(f"'questions' is not a list: {type(questions)}")
            return []
        
        # Normalize fields
        normalized = []
        for q in questions:
            if isinstance(q, dict):
                # Normalize question -> text
                if "question" in q and "text" not in q:
                    q["text"] = q.pop("question")
                # Ensure answer field exists
                if "answer" not in q:
                    q["answer"] = q.get("correct_answer", q.get("correct", q.get("model_answer", "")))
                # For MCQs, format options into text if present
                if "options" in q and isinstance(q["options"], (list, dict)):
                    opts = q["options"]
                    if isinstance(opts, dict):
                        opt_text = "\n".join([f"{k}) {v}" for k, v in opts.items()])
                    else:
                        opt_text = "\n".join([f"{chr(65+i)}) {o}" for i, o in enumerate(opts)])
                    q["text"] = q.get("text", "") + "\n" + opt_text
                normalized.append(q)
        
        logger.info(f"Parsed {len(normalized)} questions successfully")
        return normalized
    except json.JSONDecodeError as e:
        logger.error(f"JSON decode error: {e}. Response: {response_text[:500]}")
        return []
    except Exception as e:
        logger.error(f"Parse error: {e}")
        return []

async def run_batch_query(prompt: PromptTemplate, q_type: str, num: int, marks: int, context: str, subject: str, difficulty: str) -> List[Dict[str, str]]:
    if num <= 0: return []
    try:
        logger.info(f"Generating {num} {q_type} questions...")
        chain = prompt | llm | StrOutputParser()
        response = await chain.ainvoke({
            "subject": subject, "num_questions": num, "question_type": q_type,
            "difficulty": difficulty, "marks": marks, "context": context
        })
        logger.info(f"Raw LLM Response for {q_type}: {response[:500]}...") # Log first 500 chars
        
        parsed = parse_json_output(response)
        if not parsed:
            logger.error(f"Failed to parse {q_type} response. Full response saved to debug log.")
            logger.info(f"FULL FAILED RESPONSE ({q_type}): {response}")
            
        return parsed
    except Exception as e:
        logger.error(f"Error in batch query for {q_type}: {e}")
        return []

async def generate_question_paper(docs_content: tuple, subject: str, pattern: str, difficulty: str) -> Dict[str, List]:
    # Generator Patterns (Separate from Evaluator)
    patterns = {
        "CIA": {"mcq": (10, 1), "short": (5, 4), "long": (2, 10)},
        "Model": {"mcq": (10, 1), "short": (5, 5), "long": (5, 8)}
    }
    config = patterns.get(pattern)
    if not config: raise ValueError(f"Invalid pattern '{pattern}'")

    num_mcq, marks_mcq = config["mcq"]
    num_short, marks_short = config["short"]
    num_long, marks_long = config["long"]
    
    context_sample = "\n---\n".join(random.sample(list(docs_content), min(len(docs_content), 5)))
    paper = {"MCQ": [], "Short": [], "Long": []}

    paper["MCQ"] = await run_batch_query(MCQ_BATCH_PROMPT, "MCQ", num_mcq, marks_mcq, context_sample, subject, difficulty)
    paper["Short"] = await run_batch_query(RUBRIC_BATCH_PROMPT, "Short Answer", num_short * 2, marks_short, context_sample, subject, difficulty)
    paper["Long"] = await run_batch_query(RUBRIC_BATCH_PROMPT, "Long Essay", num_long * 2, marks_long, context_sample, subject, difficulty)

    # Filter out non-dict items and assign marks (LLM sometimes returns malformed data)
    paper["MCQ"] = [item for item in paper["MCQ"] if isinstance(item, dict)]
    paper["Short"] = [item for item in paper["Short"] if isinstance(item, dict)]
    paper["Long"] = [item for item in paper["Long"] if isinstance(item, dict)]
    
    for item in paper["MCQ"]: item["marks"] = marks_mcq
    for item in paper["Short"]: item["marks"] = marks_short
    for item in paper["Long"]: item["marks"] = marks_long
        
    return paper

# --- DOCX UTILS ---
# --- DOCX UTILS ---
import latex2mathml.converter
from lxml import etree

try:
    # Load transformation stylesheet
    xslt_tree = etree.parse(str(BASE_DIR / "MML2OMML.xsl"))
    xslt_transform = etree.XSLT(xslt_tree)
except Exception as e:
    logger.error(f"Failed to load MML2OMML.xsl: {e}")
    xslt_transform = None

def latex_to_omml(latex_str):
    if not xslt_transform: return None
    try:
        # 1. Convert LaTeX to MathML (no XML header)
        mathml = latex2mathml.converter.convert(latex_str)
        # latex2mathml might produce <math ...> ... </math>
        
        # 2. Transform MathML to OMML
        # We need to parse MathML string to Element
        mathml_tree = etree.fromstring(mathml)
        omml_tree = xslt_transform(mathml_tree)
        
        return omml_tree
    except Exception as e:
        logger.error(f"Math conversion failed for '{latex_str}': {e}")
        return None

def replace_text_in_paragraph(paragraph: Paragraph, context: Dict[str, str]):
    full_text = "".join(run.text for run in paragraph.runs)
    if '{{' not in full_text: return
    
    # Perform Replacement
    for key, value in context.items():
        if key in full_text: full_text = full_text.replace(key, str(value))
    
    # Check if there is any LaTeX to render ($...$)
    # Logic: 
    # 1. Start fresh paragraph
    # 2. Split full_text by regex matches
    # 3. Add runs for text, add OMML for math
    
    # Regex for $...$ (non-greedy)
    # Note: We assume single dollar signs are used as per prompt instructions
    segments = re.split(r'(\$.*?\$)', full_text)
    
    # Clear existing runs
    p = paragraph._p
    for run in paragraph.runs:
        p.remove(run._r)
        
    for segment in segments:
        if not segment: continue
        
        # Check if math
        if segment.startswith('$') and segment.endswith('$') and len(segment) > 2:
            latex_content = segment[1:-1] # Strip $
            omml_element = latex_to_omml(latex_content)
            
            if omml_element is not None:
                # Insert OMML
                # Using lxml element directly might fail if python-docx expects OxmlElement?
                # Actually, python-docx uses lxml.etree internally as well.
                # However, to be safe, we convert our lxml Element to string and parse with docx
                from docx.oxml import parse_xml, OxmlElement
                from docx.oxml.ns import nsdecls
                
                # Wrap in oMathPara if needed? Usually OMML returns oMath.
                # Inline math should just be oMath.
                # Let's inspect omml_element type. It is an lxml Element.
                
                # We need to append it to paragraph._p
                # python-docx elements are lxml _Element wrappers.
                # We can just append the lxml element directly if it matches?
                
                # Safer: Serialize and re-parse with docx's parse_xml to ensure correct class casting
                omml_xml_str = etree.tostring(omml_element, encoding='unicode')
                
                # MML2OMML result usually includes namespaces.
                # Sometimes we need to strip namespaces or ensure they match.
                # But parse_xml logic usually handles it.
                try:
                    oxml_obj = parse_xml(omml_xml_str)
                    paragraph._p.append(oxml_obj)
                except Exception as ex:
                    logger.error(f"Failed to insert OMML: {ex}")
                    paragraph.add_run(segment) # Fallback
            else:
                 paragraph.add_run(segment) # Fallback text
        else:
            # Regular text
            # Handle newlines if any
            if '\n' in segment:
                lines = segment.split('\n')
                for i, line in enumerate(lines):
                    if i > 0: paragraph.add_run().add_break()
                    paragraph.add_run(line)
            else:
                paragraph.add_run(segment)

def replace_placeholders(doc: Document, context: Dict[str, str]):
    for p in doc.paragraphs: replace_text_in_paragraph(p, context)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: replace_text_in_paragraph(p, context)
    for section in doc.sections:
        for p in section.header.paragraphs: replace_text_in_paragraph(p, context)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: replace_text_in_paragraph(p, context)
        for p in section.footer.paragraphs: replace_text_in_paragraph(p, context)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: replace_text_in_paragraph(p, context)
    return doc

def parse_answer_and_marks(answer_text: str):
    if not answer_text: return "", ""
    lines = answer_text.strip().split('\n')
    cleaned_answer_lines = []
    formatted_marks_lines = []
    # Updated pattern to match (3), (2.5), (3 marks), (2.5 marks), etc.
    mark_pattern = re.compile(r'\s*\((\d+(?:\.\d+)?)(?:\s*marks?)?\)\s*$', re.IGNORECASE)

    for line in lines:
        line = line.strip()
        # Skip Keywords line from marks
        if line.startswith('**Keywords') or line.startswith('Keywords'):
            cleaned_answer_lines.append(line)
            formatted_marks_lines.append("")
            continue
            
        match = mark_pattern.search(line)
        if match:
            mark_number = match.group(1)
            formatted_marks_lines.append(f"({mark_number})")
            cleaned_line = mark_pattern.sub('', line).strip()
            cleaned_answer_lines.append(cleaned_line)
        else:
            cleaned_answer_lines.append(line)
            formatted_marks_lines.append("")

    return "\n".join(cleaned_answer_lines), "\n".join(formatted_marks_lines)

# --- ROUTES ---
@app.get("/")
async def read_root():
    return {"message": "KCLAS Question Paper Generator API", "version": "2.0"}

@app.post("/login")
async def login(request: Request, email: str = Form(...), password: str = Form(...)):
    try:
        user = await db_module.find_user_by_email(email)
        if user and bcrypt.checkpw(password.encode('utf-8'), user["password"].encode('utf-8')):
            return RedirectResponse(url="/dashboard", status_code=303)
        raise HTTPException(status_code=401, detail="Invalid credentials")
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=500, detail="Login failed")

@app.get("/departments")
async def get_departments():
    return await db_module.get_departments()

@app.get("/details/{department}")
async def get_details_by_department(department: str):
    try:
        dept_clean = department.strip()
        return await db_module.get_details_by_department(dept_clean)
    except Exception as e:
        logger.error(f"Error fetching details for {department}: {e}")
        return []

@app.get("/saved-papers")
async def get_saved_papers():
    try:
        return await db_module.get_saved_papers()
    except Exception as e:
        logger.error(f"Error fetching saved papers: {e}")
        return []

@app.get("/saved-paper/{paper_id}")
async def get_saved_paper(paper_id: str):
    try:
        paper = await db_module.get_saved_paper(int(paper_id))
        if not paper:
            raise HTTPException(status_code=404, detail="Paper not found")
        return paper
    except Exception as e:
        logger.error(f"Error fetching paper {paper_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-syllabus")
async def upload_syllabus(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    temp_file_path = TEMP_DIR / f"temp_{random.randint(1000, 9999)}_{file.filename}"
    try:
        with open(temp_file_path, "wb") as f: f.write(await file.read())
        units = extract_units_from_pdf(str(temp_file_path))
        return {"units": units}
    finally:
        if temp_file_path.exists(): temp_file_path.unlink()

@app.post("/generate-questions")
async def generate_questions(request: Request):
    try:
        data = await request.json()
        exam_type = "Model" if data.get("exam_type") == "Models" else data.get("exam_type")
        doc_chunks = create_document_chunks(data.get("selected_units", []))
        if not doc_chunks: raise HTTPException(status_code=422, detail="Failed to process content.")
        
        paper = await generate_question_paper(
            tuple(doc.page_content for doc in doc_chunks), data["subject"], exam_type, data["difficulty"]
        )
        
        # --- SAVE TO HISTORY (PostgreSQL) ---
        try:
            await db_module.insert_question_paper(
                subject=data["subject"],
                exam_type=exam_type,
                difficulty=data["difficulty"],
                paper=paper
            )
        except Exception as e:
            logger.error(f"Failed to save paper content: {e}")

        return {"question_paper": paper}
    except Exception as e:
        logger.error(f"Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/download-paper")
async def download_paper(data: DownloadRequest):
    exam_type = "Model" if data.examType == "Models" else data.examType
    template_path = TEMPLATE_PATHS.get(exam_type, {}).get("paper")
    if not template_path or not template_path.exists():
        raise HTTPException(status_code=404, detail="Template not found")

    doc = Document(template_path)
    context = {
        "{{department}}": data.department, "{{batch}}": data.batch, "{{semester}}": data.semester,
        "{{subject}}": data.subject, "{{examType}}": data.examType, "{{duration}}": data.duration,
        "{{paperSetter}}": data.paperSetter, "{{hod}}": data.hod,
    }
    mcqs = [q for q in data.questions if q.type == "MCQ"]
    short_answers = [q for q in data.questions if q.type == "Short Answer"]
    long_essays = [q for q in data.questions if q.type == "Long Essay"]

    for i in range(1, 11):
        if i <= len(mcqs):
            q = mcqs[i - 1]
            parts = q.text.split('\n')
            context[f"{{{{Q{i}}}}}"] = parts[0]
            context[f"{{{{Q{i}_A}}}}"] = parts[1] if len(parts) > 1 else ""
            context[f"{{{{Q{i}_B}}}}"] = parts[2] if len(parts) > 2 else ""
            context[f"{{{{Q{i}_C}}}}"] = parts[3] if len(parts) > 3 else ""
            context[f"{{{{Q{i}_D}}}}"] = parts[4] if len(parts) > 4 else ""
    
    q_num = 11
    for i in range(0, len(short_answers), 2):
        context[f"{{{{Q{q_num}a}}}}"] = short_answers[i].text
        context[f"{{{{Q{q_num}b}}}}"] = short_answers[i + 1].text if i + 1 < len(short_answers) else ""
        q_num += 1

    q_num = 16
    for i in range(0, len(long_essays), 2):
        context[f"{{{{Q{q_num}a}}}}"] = long_essays[i].text
        context[f"{{{{Q{q_num}b}}}}"] = long_essays[i + 1].text if i + 1 < len(long_essays) else ""
        q_num += 1

    doc = replace_placeholders(doc, context)
    file_path = TEMP_DIR / f"{data.subject}_QP_{random.randint(1000,9999)}.docx"
    doc.save(file_path)
    return FileResponse(path=file_path, filename=f"{data.subject}_Question_Paper.docx")

@app.post("/download-key")
async def download_key(data: DownloadRequest):
    exam_type = "Model" if data.examType == "Models" else data.examType
    template_path = TEMPLATE_PATHS.get(exam_type, {}).get("key")
    if not template_path or not template_path.exists():
        raise HTTPException(status_code=404, detail="Key template not found")
    
    doc = Document(template_path)
    context = {
        "{{department}}": data.department, "{{batch}}": data.batch, "{{semester}}": data.semester,
        "{{subject}}": data.subject, "{{examType}}": data.examType, "{{duration}}": data.duration,
        "{{paperSetter}}": data.paperSetter, "{{hod}}": data.hod,
    }
    mcqs = [q for q in data.questions if q.type == "MCQ"]
    short = [q for q in data.questions if q.type == "Short Answer"]
    long = [q for q in data.questions if q.type == "Long Essay"]

    for i in range(1, 11):
        context[f"{{{{A{i}}}}}"] = mcqs[i-1].answer if i <= len(mcqs) else ""

    q_num = 11
    for i in range(0, len(short), 2):
        a, m = parse_answer_and_marks(short[i].answer.replace('\\n', '\n'))
        context[f"{{{{A{q_num}a}}}}"] = a
        context[f"{{{{M{q_num}a}}}}"] = m
        if i + 1 < len(short):
            a2, m2 = parse_answer_and_marks(short[i+1].answer.replace('\\n', '\n'))
            context[f"{{{{A{q_num}b}}}}"] = a2
            context[f"{{{{M{q_num}b}}}}"] = m2
        q_num += 1

    q_num = 16
    for i in range(0, len(long), 2):
        a, m = parse_answer_and_marks(long[i].answer.replace('\\n', '\n'))
        context[f"{{{{A{q_num}a}}}}"] = a
        context[f"{{{{M{q_num}a}}}}"] = m
        if i + 1 < len(long):
            a2, m2 = parse_answer_and_marks(long[i+1].answer.replace('\\n', '\n'))
            context[f"{{{{A{q_num}b}}}}"] = a2
            context[f"{{{{M{q_num}b}}}}"] = m2
        q_num += 1
        
    doc = replace_placeholders(doc, context)
    file_path = TEMP_DIR / f"{data.subject}_Key_{random.randint(1000,9999)}.docx"
    doc.save(file_path)
    return FileResponse(path=file_path, filename=f"{data.subject}_Answer_Key.docx")

if __name__ == "__main__":
    import uvicorn
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--port", type=int, default=8000)
    args = parser.parse_args()
    uvicorn.run(app, host="0.0.0.0", port=args.port)